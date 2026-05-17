"""
Dokumenttiparser - tukee PDF- ja Word-tiedostoja.

PDF: PyMuPDF
Word (.docx): python-docx
"""
from __future__ import annotations
import logging
from dataclasses import dataclass
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """Puretun dokumentin sisältö."""
    text: str
    pages: list[str]
    page_count: int
    filename: str
    has_text_layer: bool
    metadata: dict
    file_type: str   # "pdf" tai "docx"


# Taaksepäin yhteensopivuus - vanha nimi toimii edelleen
ParsedPDF = ParsedDocument


def extract_text(file_bytes: bytes, filename: str = "document.pdf") -> ParsedDocument:
    """
    Purkaa tekstin PDF- tai Word-tiedostosta tiedostonimen perusteella.

    Args:
        file_bytes: Tiedoston sisältö tavuina
        filename: Tiedostonimi (määrittää parserin)

    Returns:
        ParsedDocument sisältäen tekstin ja metatiedot

    Raises:
        ValueError: Jos tiedosto ei ole tuettu tai kelvollinen
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(file_bytes, filename)
    elif suffix == ".docx":
        return _extract_docx(file_bytes, filename)
    else:
        raise ValueError(f"Tiedostotyyppi '{suffix}' ei ole tuettu. Tuetut: .pdf, .docx")


def _extract_pdf(pdf_bytes: bytes, filename: str) -> ParsedDocument:
    """Purkaa tekstin PDF:stä PyMuPDF:lla."""
    try:
        import fitz
    except ImportError:
        raise RuntimeError("PyMuPDF ei ole asennettu. Asenna: pip install pymupdf")

    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as e:
        raise ValueError(f"PDF:n avaaminen epäonnistui: {e}")

    page_texts: list[str] = []
    total_chars = 0

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text")
        page_texts.append(text)
        total_chars += len(text.strip())

    has_text_layer = total_chars > 50

    if not has_text_layer:
        logger.warning(f"PDF '{filename}' ei näytä sisältävän tekstikerrosta ({total_chars} merkkiä).")
        _ocr_hook_placeholder(filename)

    meta = doc.metadata or {}
    metadata = {
        "title": meta.get("title", ""),
        "author": meta.get("author", ""),
        "creator": meta.get("creator", ""),
        "creation_date": meta.get("creationDate", ""),
    }
    doc.close()

    full_text = _join_pages(page_texts)

    return ParsedDocument(
        text=full_text,
        pages=page_texts,
        page_count=len(page_texts),
        filename=filename,
        has_text_layer=has_text_layer,
        metadata=metadata,
        file_type="pdf",
    )


def _extract_docx(docx_bytes: bytes, filename: str) -> ParsedDocument:
    """Purkaa tekstin Word-dokumentista python-docx:lla."""
    try:
        import docx
        import io
    except ImportError:
        raise RuntimeError("python-docx ei ole asennettu. Asenna: pip install python-docx")

    try:
        doc = docx.Document(io.BytesIO(docx_bytes))
    except Exception as e:
        raise ValueError(f"Word-dokumentin avaaminen epäonnistui: {e}")

    # Kerää kappaleet
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    # Kerää myös taulukoiden teksti
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                paragraphs.append(" | ".join(row_texts))

    full_text = "\n".join(paragraphs)
    has_text_layer = len(full_text.strip()) > 50

    # Word-dokumenteilla ei ole sivuja samalla tavalla - käsitellään yhtenä "sivuna"
    pages = [full_text]

    # Metatiedot
    core_props = doc.core_properties
    metadata = {
        "title": core_props.title or "",
        "author": core_props.author or "",
        "creator": core_props.last_modified_by or "",
        "creation_date": str(core_props.created) if core_props.created else "",
    }

    logger.info(f"Word-dokumentti '{filename}' purettu: {len(full_text)} merkkiä")

    return ParsedDocument(
        text=full_text,
        pages=pages,
        page_count=1,
        filename=filename,
        has_text_layer=has_text_layer,
        metadata=metadata,
        file_type="docx",
    )


def extract_text_from_path(path: Path) -> ParsedDocument:
    """Apufunktio tiedostopolusta lukemiseen (testaukseen)."""
    file_bytes = path.read_bytes()
    return extract_text(file_bytes, filename=path.name)


def _join_pages(pages: list[str]) -> str:
    """Yhdistää sivut selkeästi eroteltuina."""
    parts = []
    for i, page_text in enumerate(pages, start=1):
        if page_text.strip():
            parts.append(f"[Sivu {i}]\n{page_text.strip()}")
    return "\n\n".join(parts)


def _ocr_hook_placeholder(filename: str) -> None:
    """OCR-hook - tuleva ominaisuus."""
    logger.info(
        f"OCR-hook aktivoitu tiedostolle '{filename}'. "
        "OCR-tunnistus ei ole vielä käytössä."
    )
